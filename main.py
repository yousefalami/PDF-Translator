# main.py
# -*- coding: utf-8 -*-
import os
import configparser
import requests
import pdfplumber
import json
import time
import concurrent.futures
import re
import threading
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tqdm import tqdm

# --- Initial Configuration ---
CONFIG_FILENAME = 'config.ini'
DEFAULT_FONT = "Arial"
input_lock = threading.Lock()

# --- Helper Functions ---
def print_separator(char="=", length=70):
    print(f"\n{char * length}\n")

def print_header(title):
    print_separator()
    print(f"--- {title} ---")
    print()

def read_config(filename=CONFIG_FILENAME):
    if not os.path.exists(filename):
        print(f"❌ Error: Config file '{filename}' not found.")
        return None
    config = configparser.ConfigParser()
    config.read(filename)
    try:
        api = config['API']
        settings = config['SETTINGS']
        return {
            'api_url': api.get('url'),
            'api_model': api.get('model'),
            'font_name': api.get('font_name', DEFAULT_FONT),
            'prompt_template': api.get('prompt_template'),
            'max_workers': settings.getint('max_concurrent_workers', 3),
            'max_chars_batch': settings.getint('max_chars_per_batch', 12000),
            'max_retries': settings.getint('max_retries', 3),
            'retry_delay': settings.getint('initial_retry_delay', 2),
            'source_language': "English",
            'target_language': "Farsi"
        }
    except Exception as e:
        print(f"❌ Error reading config: {e}")
        return None

# --- Critical Function: Apply RTL & Font for Windows/Mac ---
def apply_rtl_formatting(paragraph, run, font_name):
    """
    Injects XML tags to enforce Right-to-Left text direction and correct
    font rendering on Windows and Mac (Word).
    """
    # 1. Set Paragraph Direction (BiDi)
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

    # 2. Set Run Direction (RTL)
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

    # 3. Set Complex Script (CS) Fonts
    # This tells Word to use the specified font (e.g., B Nazanin) for Persian text.
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman') 
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), font_name)            
    rPr.append(rFonts)

# --- PDF Extraction ---
def extract_text_from_pdf(pdf_path, start_page, end_page):
    if not os.path.exists(pdf_path):
        print(f"❌ Error: PDF file not found at '{pdf_path}'.")
        return None
    extracted_data = [] 
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            final_start = max(1, start_page)
            final_end = min(total_pages, end_page) if end_page else total_pages
            print(f"📄 Processing pages {final_start} to {final_end} (Total: {total_pages})")
            pages_to_process = pdf.pages[final_start-1 : final_end]
            for i, page in enumerate(tqdm(pages_to_process, desc="🔍 Extracting", unit="page")):
                text = page.extract_text()
                extracted_data.append((final_start + i, text if text else ""))
        return extracted_data
    except Exception as e:
        print(f"❌ Error extraction: {e}")
        return None

# --- Dynamic Batching ---
def create_dynamic_batches(extracted_data, max_chars):
    batches = []
    current_batch_items = []
    current_chars = 0
    previous_context_text = ""
    last_page_text_of_current_batch = ""

    for real_page_num, text in extracted_data:
        page_obj = {"page_id": real_page_num, "text_to_translate": text.strip()}
        obj_char_count = len(json.dumps(page_obj, ensure_ascii=False))

        if current_batch_items and (current_chars + obj_char_count > max_chars):
            batches.append({'items': current_batch_items, 'context': previous_context_text})
            previous_context_text = last_page_text_of_current_batch
            current_batch_items = []
            current_chars = 0
        
        current_batch_items.append(page_obj)
        current_chars += obj_char_count
        last_page_text_of_current_batch = text.strip()

    if current_batch_items:
        batches.append({'items': current_batch_items, 'context': previous_context_text})
    return batches

# --- API Translation ---
def translate_batch_via_api(batch_data, config):
    batch_items = batch_data['items']
    context_text = batch_data['context']
    page_ids = [item['page_id'] for item in batch_items]
    headers = {'Content-Type': 'application/json'}
    
    if len(context_text) > 2000: context_text = "..." + context_text[-2000:]
    if not context_text: context_text = "None (Start)"

    final_prompt = config['prompt_template'].format(
        source_language=config['source_language'],
        target_language=config['target_language'],
        context=context_text,
        json_data=json.dumps(batch_items, ensure_ascii=False)
    )
    payload = {"model": config['api_model'], "messages": [{"role": "user", "content": final_prompt}]}

    attempt = 0
    while True:
        try:
            if attempt > 0: print(f"⏳ Retry {attempt} for {page_ids}...")
            response = requests.post(config['api_url'], headers=headers, json=payload, timeout=300)
            if response.status_code != 200: raise Exception(f"HTTP {response.status_code}")
            
            content = response.json()["choices"][0]["message"]["content"]
            
            # Regex Strategy 1: Look for "text_to_translate"
            pattern = re.compile(r'"page_id"\s*:\s*(\d+)\s*,\s*"text_to_translate"\s*:\s*"((?:\\.|[^"\\])*)"', re.DOTALL)
            matches = pattern.findall(content)
            
            # Regex Strategy 2: Look for "translated_text" (fallback)
            if not matches:
                pattern = re.compile(r'"page_id"\s*:\s*(\d+)\s*,\s*"(?:translated_text|translation)"\s*:\s*"((?:\\.|[^"\\])*)"', re.DOTALL)
                matches = pattern.findall(content)

            if not matches: 
                print(f"⚠️ Response Error Snippet: {content[:200]}...")
                raise ValueError("No valid JSON found")

            result_dict = {}
            for match in matches:
                clean_text = match[1].replace('\\n', '\n').replace('\\"', '"').replace('\\\\', '\\')
                result_dict[int(match[0])] = clean_text
            return result_dict

        except Exception as e:
            attempt += 1
            print(f"\n❌ Error {page_ids}: {e}")
            if attempt > config['max_retries']:
                with input_lock:
                    print(f"🔴 Issue with pages {page_ids}")
                    print("Options: [Enter] Retry, [S] Skip, [Q] Quit")
                    choice = input("👉 Choice: ").strip().lower()
                    if choice == 's': return {}
                    elif choice == 'q': os._exit(1)
                    else: attempt = 0
            else:
                time.sleep(config['retry_delay'])

# --- Document Generation ---
def create_translation_document(original_pdf_name, page_data, config):
    output_folder="translated_documents"
    if not os.path.exists(output_folder): os.makedirs(output_folder)

    doc = Document()
    doc.add_heading(f"Translation: {original_pdf_name}", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    print(f"\n📝 Creating Word (Win/Mac/Mobile Compatible)...")
    for i, data in enumerate(tqdm(page_data, desc="✒️ Writing", unit="page")):
        page_num = data['page_id']
        doc.add_heading(f"Page {page_num}", level=2)
        table = doc.add_table(rows=1, cols=2, style='Table Grid')
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.5)
        
        # Left Cell: English Source
        p1 = table.rows[0].cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.add_run(data['original'] or "").font.size = Inches(0.11)
        
        # Right Cell: Farsi Translation
        p2 = table.rows[0].cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
        
        translated_text = data['translated'] or ""
        run = p2.add_run(translated_text)
        
        run.font.name = config['font_name']
        run.font.size = Pt(13)
        run.bold = False 
        
        # Apply strict RTL formatting
        apply_rtl_formatting(p2, run, config['font_name'])

        if i < len(page_data) - 1: doc.add_page_break()

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    fname = os.path.join(output_folder, f"{original_pdf_name}_{timestamp}.docx")
    doc.save(fname)
    print(f"✅ Saved: {os.path.abspath(fname)}")

# --- Main Execution ---
def main():
    print_header("Ultimate PDF Translator: Final Version")
    config = read_config()
    if not config: return

    pdf_path = input("➡️  PDF Path: ").strip().strip('"')
    default_model = "gemini-3.0-pro"
    model_in = input(f"➡️  Model (Default: {default_model}): ").strip()
    config['api_model'] = model_in if model_in else default_model

    s_page = input("➡️  Start Page (1): ").strip()
    start_page = int(s_page) if s_page.isdigit() else 1
    e_page = input("➡️  End Page (All): ").strip()
    end_page = int(e_page) if e_page.isdigit() else None

    config['source_language'] = input("➡️  Source (English): ").strip() or "English"
    config['target_language'] = input("➡️  Target (Farsi): ").strip() or "Farsi"

    extracted = extract_text_from_pdf(pdf_path, start_page, end_page)
    if not extracted: return

    batches = create_dynamic_batches(extracted, config['max_chars_batch'])
    
    print_header(f"Translating {len(batches)} batches...")
    results = {}
    with concurrent.futures.ThreadPoolExecutor(max_workers=config['max_workers']) as executor:
        futures = {executor.submit(translate_batch_via_api, b, config): b for b in batches}
        for f in tqdm(concurrent.futures.as_completed(futures), total=len(batches)):
            try: results.update(f.result())
            except Exception as e: print(f"Error: {e}")

    final_data = [{'page_id': p, 'original': t, 'translated': results.get(p, "FAILED")} for p, t in extracted]
    
    suffix = f"_p{start_page}-{end_page if end_page else 'end'}"
    create_translation_document(os.path.splitext(os.path.basename(pdf_path))[0] + suffix, final_data, config)
    
    print("\n🎉 Done. Press Enter to exit.")
    input()

if __name__ == "__main__":
    main()